#!/usr/bin/env python3
import argparse
import hashlib
import json
import os
import re
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

from docx import Document


CJK_RE = re.compile(r"[\u3400-\u9fff]")
SKIP_RE = re.compile(r"^\s*(https?://|/|\.?/|[A-Za-z]:\\|pip |conda |sudo |cd |python |python3 |git |wget |curl |npm |yarn |ros|ls\b|mkdir\b|rm\b)")


SYSTEM = """You are a professional technical translator translating zh-CN lab manuals into en-US.
Translate only the Chinese instructional prose into natural, clear American English for students.
Preserve code, commands, paths, URLs, variables, placeholders, version numbers, model names, product names, figure/table numbers, numbering, and line breaks.
Do not add explanations. Return only valid JSON in the requested shape."""


def has_cjk(text: str) -> bool:
    return bool(CJK_RE.search(text))


def should_translate(text: str) -> bool:
    stripped = text.strip()
    if not stripped or not has_cjk(stripped):
        return False
    if SKIP_RE.search(stripped):
        return False
    return True


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for row2 in nested.rows:
                        for cell2 in row2.cells:
                            for p in cell2.paragraphs:
                                yield p
    for section in doc.sections:
        parts = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for part in parts:
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def cache_key(text: str, model: str) -> str:
    return hashlib.sha256((model + "\0" + text).encode("utf-8")).hexdigest()


def call_openai(batch, model: str, timeout: int, retries: int, base_url: str, api_key_env: str):
    api_key = os.environ.get(api_key_env)
    if not api_key:
        raise RuntimeError(f"{api_key_env} is not set")
    url = base_url.rstrip("/") + "/chat/completions"
    user = {
        "instruction": "Translate each item. Return JSON exactly as {\"translations\":[{\"id\":0,\"text\":\"...\"}, ...]}. Preserve every id exactly once.",
        "items": [{"id": i, "text": text} for i, text in enumerate(batch)],
    }
    payload = {
        "model": model,
        "temperature": 0,
        "max_tokens": 8192,
        "response_format": {"type": "json_object"},
        "messages": [
            {"role": "system", "content": SYSTEM},
            {"role": "user", "content": json.dumps(user, ensure_ascii=False)},
        ],
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    last_error = None
    for attempt in range(retries + 1):
        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                body = json.loads(resp.read().decode("utf-8"))
            content = body["choices"][0]["message"]["content"]
            parsed = json.loads(content)
            raw_translations = parsed["translations"]
            if len(raw_translations) != len(batch):
                raise RuntimeError(f"translation count mismatch: {len(raw_translations)} != {len(batch)}")
            if raw_translations and isinstance(raw_translations[0], dict):
                by_id = {int(item["id"]): str(item["text"]) for item in raw_translations}
                if sorted(by_id) != list(range(len(batch))):
                    raise RuntimeError(f"translation ids mismatch: {sorted(by_id)} != {list(range(len(batch)))}")
                return [by_id[i] for i in range(len(batch))]
            return raw_translations
        except urllib.error.HTTPError as exc:
            try:
                detail = exc.read().decode("utf-8", errors="replace")[:1000]
            except Exception:
                detail = ""
            last_error = f"{exc} {detail}"
            sleep = min(180, 30 * (attempt + 1)) if exc.code == 429 else min(60, 2 ** attempt)
            print(f"OpenAI call failed on attempt {attempt + 1}/{retries + 1}: {last_error}; sleeping {sleep}s", file=sys.stderr)
            time.sleep(sleep)
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, KeyError, IndexError, RuntimeError) as exc:
            last_error = exc
            sleep = min(60, 2 ** attempt)
            print(f"OpenAI call failed on attempt {attempt + 1}/{retries + 1}: {exc}; sleeping {sleep}s", file=sys.stderr)
            time.sleep(sleep)
    raise RuntimeError(f"OpenAI call failed after retries: {last_error}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--cache", required=True)
    parser.add_argument("--model", default=os.environ.get("TRANSLATION_MODEL", "gpt-4.1-mini"))
    parser.add_argument("--base-url", default=os.environ.get("TRANSLATION_BASE_URL", "https://api.openai.com/v1"))
    parser.add_argument("--api-key-env", default=os.environ.get("TRANSLATION_API_KEY_ENV", "OPENAI_API_KEY"))
    parser.add_argument("--batch-chars", type=int, default=4500)
    parser.add_argument("--batch-items", type=int, default=12)
    parser.add_argument("--timeout", type=int, default=120)
    parser.add_argument("--retries", type=int, default=5)
    parser.add_argument("--sleep-between", type=float, default=0)
    parser.add_argument("--limit", type=int, default=0, help="Translate at most this many uncached paragraphs; 0 means all.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    cache_path = Path(args.cache)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    cache_path.parent.mkdir(parents=True, exist_ok=True)

    cache = {}
    if cache_path.exists():
        cache = json.loads(cache_path.read_text(encoding="utf-8"))

    doc = Document(input_path)
    paragraphs = list(iter_paragraphs(doc))
    todo = []
    for idx, p in enumerate(paragraphs):
        text = p.text
        if should_translate(text):
            key = cache_key(text, args.model)
            if key in cache:
                set_paragraph_text(p, cache[key])
            else:
                todo.append((idx, text, key))

    if args.limit:
        todo = todo[: args.limit]

    print(f"paragraphs={len(paragraphs)} uncached_to_translate={len(todo)} model={args.model} base_url={args.base_url} key_env={args.api_key_env}")
    batch = []
    batch_meta = []
    batch_chars = 0
    translated_count = 0

    def translate_batch(items, metas):
        if not items:
            return
        try:
            translations = call_openai(items, args.model, args.timeout, args.retries, args.base_url, args.api_key_env)
        except RuntimeError:
            if len(items) == 1:
                raise
            mid = len(items) // 2
            translate_batch(items[:mid], metas[:mid])
            translate_batch(items[mid:], metas[mid:])
            return
        for (idx, _text, key), translated in zip(metas, translations):
            translated = str(translated).strip()
            cache[key] = translated
            set_paragraph_text(paragraphs[idx], translated)

    def flush():
        nonlocal batch, batch_meta, batch_chars, translated_count
        if not batch:
            return
        before = len(cache)
        translate_batch(batch, batch_meta)
        if args.sleep_between:
            time.sleep(args.sleep_between)
        translated_count += len(batch_meta)
        cache_path.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
        doc.save(output_path)
        print(f"translated={translated_count}/{len(todo)} cache={len(cache)} new_cache={len(cache)-before} output={output_path}")
        batch = []
        batch_meta = []
        batch_chars = 0

    for meta in todo:
        _idx, text, _key = meta
        if batch and (batch_chars + len(text) > args.batch_chars or len(batch) >= args.batch_items):
            flush()
        batch.append(text)
        batch_meta.append(meta)
        batch_chars += len(text)
    flush()
    doc.save(output_path)
    print(f"done output={output_path}")


if __name__ == "__main__":
    main()
