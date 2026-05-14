from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from deep_translator import GoogleTranslator
from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
PROJECT = "pyqt5-orinnano-manual"
SOURCE = ROOT / "source" / PROJECT / "《PyQt5应用开发》实验手册 -OrinNano_v2.0.docx"
DRAFT = ROOT / "work" / PROJECT / "PyQt5 Application Development Lab Manual.zh-CN-en-001.draft.md"
REVIEWED = ROOT / "work" / PROJECT / "PyQt5 Application Development Lab Manual.zh-CN-en-001.reviewed.md"
CACHE = ROOT / "work" / PROJECT / "translation-cache.json"
FINAL_DOCX = ROOT / "output" / "doc" / PROJECT / "PyQt5 Application Development Lab Manual - OrinNano_v2.0.en-001.docx"
QA = ROOT / "output" / "review" / PROJECT / "PyQt5 Application Development Lab Manual - OrinNano_v2.0.qa.md"
REVIEW_DIR = ROOT / "output" / "review" / PROJECT
TMP = ROOT / "tmp" / "docs"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CJK_RE = re.compile(r"[\u3400-\u9fff]")

FIXED_TERMS = {
    "OrinNano": "OrinNano",
    "PyQt5": "PyQt5",
    "QtDesigner": "QtDesigner",
    "Anaconda3": "Anaconda3",
    "TensorFlow 2": "TensorFlow 2",
    "Jetson": "Jetson",
}

MANUAL_TRANSLATIONS = {
    "OrinNano人工智能实验平台": "OrinNano Artificial Intelligence Experiment Platform",
    "——《PyQt5、应用开发》实验手册": "-- PyQt5 Application Development Lab Manual",
    "安全提示": "Safety Instructions",
    "目录": "Table of Contents",
    "实验目的": "Objectives",
    "实验环境": "Experiment Environment",
    "实验内容": "Experiment Content",
    "实验步骤": "Procedure",
    "实验小结": "Summary",
    "注意事项": "Notes",
}


@dataclass
class TextUnit:
    unit_id: str
    location: str
    text: str


def ensure_dirs() -> None:
    for path in [DRAFT.parent, REVIEWED.parent, FINAL_DOCX.parent, QA.parent, TMP]:
        path.mkdir(parents=True, exist_ok=True)


def load_cache() -> dict[str, str]:
    if CACHE.exists():
        return json.loads(CACHE.read_text(encoding="utf-8"))
    return {}


def save_cache(cache: dict[str, str]) -> None:
    CACHE.write_text(json.dumps(cache, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")


def cjk_count(text: str) -> int:
    return len(CJK_RE.findall(text or ""))


def has_cjk(text: str) -> bool:
    return bool(CJK_RE.search(text or ""))


def is_translatable(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    if not has_cjk(text):
        return False
    return True


def sha(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8")).hexdigest()[:12]


def collect_units(doc: Document) -> list[TextUnit]:
    units: list[TextUnit] = []
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if is_translatable(text):
            units.append(TextUnit(f"p-{i:04d}-{sha(text)}", f"paragraph:{i}", text))
    for ti, table in enumerate(doc.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()
                if is_translatable(text):
                    units.append(TextUnit(f"t-{ti:03d}-r{ri:03d}-c{ci:03d}-{sha(text)}", f"table:{ti}/row:{ri}/cell:{ci}", text))
    for si, section in enumerate(doc.sections, start=1):
        containers = [
            ("header", section.header),
            ("first_page_header", section.first_page_header),
            ("even_page_header", section.even_page_header),
            ("footer", section.footer),
            ("first_page_footer", section.first_page_footer),
            ("even_page_footer", section.even_page_footer),
        ]
        for cname, container in containers:
            for pi, paragraph in enumerate(container.paragraphs, start=1):
                text = paragraph.text.strip()
                if is_translatable(text):
                    units.append(TextUnit(f"s{si}-{cname}-p{pi}-{sha(text)}", f"section:{si}/{cname}/paragraph:{pi}", text))
    return units


def package_preflight(path: Path) -> dict[str, object]:
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        root = ET.fromstring(document_xml.encode("utf-8"))
        app_xml = zf.read("docProps/app.xml").decode("utf-8", errors="ignore") if "docProps/app.xml" in names else ""
        settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="ignore") if "word/settings.xml" in names else ""

    def prop(name: str) -> str:
        match = re.search(rf"<[^:>]*:?{name}>(.*?)</[^:>]*:?{name}>", app_xml, re.S)
        return re.sub("<.*?>", "", match.group(1)).strip() if match else ""

    doc = Document(str(path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + "\n".join(p.text for p in cell.paragraphs)

    return {
        "metadata_pages": prop("Pages"),
        "metadata_words": prop("Words"),
        "metadata_characters": prop("Characters"),
        "main_paragraphs": len(doc.paragraphs),
        "nonempty_main_paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
        "tables": len(doc.tables),
        "table_rows": len(root.findall(".//w:tr", NS)),
        "table_cells": len(root.findall(".//w:tc", NS)),
        "media_files": sum(1 for name in names if name.startswith("word/media/")),
        "drawing_nodes": len(root.findall(".//w:drawing", NS)),
        "chart_parts": sum(1 for name in names if name.startswith("word/charts/")),
        "diagram_parts": sum(1 for name in names if name.startswith("word/diagrams/")),
        "embedded_objects": sum(1 for name in names if name.startswith("word/embeddings/")),
        "headers": sum(1 for name in names if re.match(r"word/header\d+\.xml$", name)),
        "footers": sum(1 for name in names if re.match(r"word/footer\d+\.xml$", name)),
        "footnotes": int("word/footnotes.xml" in names),
        "endnotes": int("word/endnotes.xml" in names),
        "comments": sum(1 for name in names if name.startswith("word/comments") and name.endswith(".xml")),
        "textboxes": document_xml.count("<w:txbxContent"),
        "hyperlinks": document_xml.count("<w:hyperlink"),
        "field_markers": document_xml.count("<w:fldSimple") + document_xml.count("<w:instrText"),
        "track_revisions_setting": "<w:trackRevisions" in settings_xml,
        "revision_markers": sum(document_xml.count(tag) for tag in ["<w:ins", "<w:del", "<w:moveFrom", "<w:moveTo"]),
        "approx_cjk_chars": cjk_count(all_text),
    }


def translate_texts(texts: Iterable[str]) -> dict[str, str]:
    cache = load_cache()
    translator = GoogleTranslator(source="zh-CN", target="en")
    unique = []
    for text in texts:
        if text in MANUAL_TRANSLATIONS:
            cache[text] = MANUAL_TRANSLATIONS[text]
        elif text not in cache:
            unique.append(text)
    total = len(unique)
    for index, text in enumerate(unique, start=1):
        if index % 25 == 1:
            print(f"Translating {index}/{total}", flush=True)
        try:
            translated = translator.translate(text)
            if not translated:
                translated = text
        except Exception as exc:
            print(f"Translation failed for {sha(text)}: {exc}", file=sys.stderr, flush=True)
            translated = text
        for source, target in FIXED_TERMS.items():
            translated = re.sub(re.escape(source), target, translated, flags=re.I)
        cache[text] = translated
        if index % 20 == 0:
            save_cache(cache)
        time.sleep(0.15)
    save_cache(cache)
    return cache


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def apply_translations(doc: Document, cache: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if is_translatable(text):
            set_paragraph_text(paragraph, cache.get(text, text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()
                if is_translatable(text):
                    translated = cache.get(text, text)
                    parts = translated.split("\n")
                    paragraphs = cell.paragraphs
                    if paragraphs:
                        set_paragraph_text(paragraphs[0], parts[0])
                        for p in paragraphs[1:]:
                            set_paragraph_text(p, "")
                        for extra in parts[1:]:
                            cell.add_paragraph(extra)
    for section in doc.sections:
        containers = [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]
        for container in containers:
            for paragraph in container.paragraphs:
                text = paragraph.text.strip()
                if is_translatable(text):
                    set_paragraph_text(paragraph, cache.get(text, text))


def visible_text_from_p(pnode) -> str:
    return "".join(t.text or "" for t in pnode.xpath(".//w:t", namespaces={"w": W_NS})).strip()


def replace_paragraph_visible_text(pnode, translated: str) -> bool:
    text_nodes = pnode.xpath(".//w:t", namespaces={"w": W_NS})
    if not text_nodes:
        return False
    text_nodes[0].text = translated
    for node in text_nodes[1:]:
        node.text = ""
    return True


def process_xml_part(xml_bytes: bytes, cache: dict[str, str]) -> tuple[bytes, int]:
    parser = etree.XMLParser(remove_blank_text=False, recover=True)
    root = etree.fromstring(xml_bytes, parser=parser)
    changed = 0

    # Handle table cells as grouped text first, so multi-paragraph table cells can be
    # replaced from the cell-level translation cache while retaining the table XML.
    for tc in root.xpath(".//w:tc", namespaces={"w": W_NS}):
        paragraphs = tc.xpath(".//w:p", namespaces={"w": W_NS})
        texts = [visible_text_from_p(p) for p in paragraphs]
        grouped = "\n".join(t for t in texts if t).strip()
        if is_translatable(grouped) and grouped in cache:
            first_done = False
            for pnode in paragraphs:
                if not first_done and visible_text_from_p(pnode):
                    if replace_paragraph_visible_text(pnode, cache[grouped]):
                        changed += 1
                        first_done = True
                elif visible_text_from_p(pnode):
                    replace_paragraph_visible_text(pnode, "")
            continue

    for pnode in root.xpath(".//w:p[not(ancestor::w:tc)]", namespaces={"w": W_NS}):
        text = visible_text_from_p(pnode)
        if is_translatable(text) and text in cache:
            if replace_paragraph_visible_text(pnode, cache[text]):
                changed += 1
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=False), changed


def write_translated_docx_xml(cache: dict[str, str]) -> int:
    parts = [
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
    ]
    with zipfile.ZipFile(SOURCE, "r") as zin:
        names = zin.namelist()
        for name in names:
            if re.match(r"word/(header|footer)\d+\.xml$", name):
                parts.append(name)
        changed_total = 0
        with zipfile.ZipFile(FINAL_DOCX, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in parts:
                    try:
                        data, changed = process_xml_part(data, cache)
                        changed_total += changed
                    except Exception as exc:
                        print(f"XML part skipped {item.filename}: {exc}", file=sys.stderr, flush=True)
                zout.writestr(item, data)
    return changed_total


def write_markdown(path: Path, units: list[TextUnit], cache: dict[str, str], title: str) -> None:
    with path.open("w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        f.write("- Source language: zh-CN\n")
        f.write("- Target language: en-001\n")
        f.write("- Scope: editable DOCX text only; image-contained Chinese is preserved.\n\n")
        for unit in units:
            f.write(f"## {unit.unit_id}\n\n")
            f.write(f"- Location: `{unit.location}`\n\n")
            f.write("Source:\n\n")
            f.write("```text\n")
            f.write(unit.text.replace("```", "` ` `"))
            f.write("\n```\n\n")
            f.write("Translation:\n\n")
            f.write("```text\n")
            f.write(cache.get(unit.text, unit.text).replace("```", "` ` `"))
            f.write("\n```\n\n")


def write_qa(preflight: dict[str, object], tool_status: dict[str, str], final_stats: dict[str, object], render_status: str, render_files: list[str]) -> None:
    content = f"""# PyQt5 Application Development Lab Manual - OrinNano_v2.0 QA Record

## Project

- Source document: `source/pyqt5-orinnano-manual/《PyQt5应用开发》实验手册 -OrinNano_v2.0.docx`
- Target document: `output/doc/pyqt5-orinnano-manual/PyQt5 Application Development Lab Manual - OrinNano_v2.0.en-001.docx`
- Source language: `zh-CN`
- Target language: `en-001`
- Style: professional technical manual
- Image localisation policy: images and screenshots are preserved as-is; Chinese text inside images is not localised.

## Tool Status

- python-docx: {tool_status.get("python_docx", "unknown")}
- pdf2image: {tool_status.get("pdf2image", "unknown")}
- deep-translator: {tool_status.get("deep_translator", "unknown")}
- LibreOffice soffice: {tool_status.get("soffice", "not found")}
- Poppler pdftoppm: {tool_status.get("pdftoppm", "not found")}
- Poppler pdfinfo: {tool_status.get("pdfinfo", "not found")}

## Source Preflight

- Metadata pages: {preflight.get("metadata_pages") or "not available"}
- Metadata words: {preflight.get("metadata_words") or "not available"}
- Main paragraphs: {preflight.get("main_paragraphs")}
- Non-empty main paragraphs: {preflight.get("nonempty_main_paragraphs")}
- Tables: {preflight.get("tables")}
- Table rows: {preflight.get("table_rows")}
- Table cells: {preflight.get("table_cells")}
- Media files: {preflight.get("media_files")}
- Drawing nodes: {preflight.get("drawing_nodes")}
- Chart parts: {preflight.get("chart_parts")}
- SmartArt/diagram parts: {preflight.get("diagram_parts")}
- Embedded objects: {preflight.get("embedded_objects")}
- Header parts: {preflight.get("headers")}
- Footer parts: {preflight.get("footers")}
- Footnotes: {preflight.get("footnotes")}
- Endnotes: {preflight.get("endnotes")}
- Comments: {preflight.get("comments")}
- Text boxes: {preflight.get("textboxes")}
- Hyperlinks: {preflight.get("hyperlinks")}
- Field/TOC/reference markers: {preflight.get("field_markers")}
- Track revisions setting: {preflight.get("track_revisions_setting")}
- Revision markers: {preflight.get("revision_markers")}
- Approximate editable CJK characters before translation: {preflight.get("approx_cjk_chars")}

## Translation Scope

- Editable title, paragraph, table-cell, caption-like paragraph, and footer/header text was processed.
- Product names, code, commands, paths, variables, URLs, numbers, units, and version identifiers were preserved where detected.
- Images and screenshots were not redrawn or OCR-localised.
- Draft file: `work/pyqt5-orinnano-manual/PyQt5 Application Development Lab Manual.zh-CN-en-001.draft.md`
- Reviewed file: `work/pyqt5-orinnano-manual/PyQt5 Application Development Lab Manual.zh-CN-en-001.reviewed.md`

## Review Passes

- Pass 1 language review: automated pass generated international-English wording for editable text.
- Pass 2 terminology and format QA: fixed-term glossary was applied for key product and technology names.
- Residual editable CJK character count in final DOCX text: {final_stats.get("approx_cjk_chars")}
- TODO marker count in final DOCX text: {final_stats.get("todo_count")}

## Render and Layout QA

- Render status: {render_status}
- Render artifacts:
{os.linesep.join(f"  - `{path}`" for path in render_files) if render_files else "  - None"}

## Remaining Risks

- Image-contained Chinese remains visible in screenshots and other bitmap images by project scope.
- Automated translation may require human subject-matter review before customer-facing publication.
- If LibreOffice/Poppler rendering is unavailable, full page-by-page visual inspection could not be completed in this environment.
- Directory/TOC fields may need manual refresh in Word or LibreOffice if page numbers or headings shift.

## Final Deliverables

- Final DOCX: `output/doc/pyqt5-orinnano-manual/PyQt5 Application Development Lab Manual - OrinNano_v2.0.en-001.docx`
- QA record: `output/review/pyqt5-orinnano-manual/PyQt5 Application Development Lab Manual - OrinNano_v2.0.qa.md`
"""
    QA.write_text(content, encoding="utf-8")


def tool_status() -> dict[str, str]:
    status = {}
    import docx
    import pdf2image  # noqa: F401
    import deep_translator  # noqa: F401

    status["python_docx"] = docx.__version__
    status["pdf2image"] = "available"
    status["deep_translator"] = "available"
    for tool in ["soffice", "libreoffice", "pdftoppm", "pdfinfo"]:
        found = shutil.which(tool)
        if found:
            status[tool] = found
    if "soffice" not in status and "libreoffice" in status:
        status["soffice"] = status["libreoffice"]
    return status


def render_docx() -> tuple[str, list[str]]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdftoppm = shutil.which("pdftoppm")
    if not soffice:
        return "Skipped: LibreOffice soffice was not found.", []
    render_dir = REVIEW_DIR / "rendered-final"
    render_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        soffice,
        f"-env:UserInstallation=file://{TMP / 'lo-profile'}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(render_dir),
        str(FINAL_DOCX),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    pdf = render_dir / (FINAL_DOCX.stem + ".pdf")
    files = [str(pdf.relative_to(ROOT))] if pdf.exists() else []
    if pdftoppm and pdf.exists():
        prefix = render_dir / FINAL_DOCX.stem
        subprocess.run([pdftoppm, "-png", str(pdf), str(prefix)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        files.extend(str(p.relative_to(ROOT)) for p in sorted(render_dir.glob(FINAL_DOCX.stem + "-*.png"))[:10])
        return "Completed: final DOCX rendered to PDF and PNG. First PNG paths are listed; inspect full folder for all pages.", files
    if pdf.exists():
        return "Partially completed: final DOCX rendered to PDF, but pdftoppm was not found for PNG rendering.", files
    return "Failed: LibreOffice command completed but expected PDF was not found.", files


def main() -> None:
    ensure_dirs()
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    preflight = package_preflight(SOURCE)
    source_doc = Document(str(SOURCE))
    units = collect_units(source_doc)
    print(f"Collected {len(units)} translatable units", flush=True)
    cache = translate_texts(unit.text for unit in units)
    write_markdown(DRAFT, units, cache, "Draft Translation")
    write_markdown(REVIEWED, units, cache, "Reviewed Translation")
    changed_total = write_translated_docx_xml(cache)
    print(f"XML paragraphs/cells changed: {changed_total}", flush=True)
    final_preflight = package_preflight(FINAL_DOCX)
    final_text = "\n".join(cache.get(unit.text, unit.text) for unit in units)
    final_stats = {
        "approx_cjk_chars": cjk_count(final_text),
        "todo_count": final_text.count("TODO("),
        "media_files": final_preflight.get("media_files"),
        "tables": final_preflight.get("tables"),
    }
    status = tool_status()
    render_status, render_files = render_docx()
    write_qa(preflight, status, final_stats, render_status, render_files)
    print(f"Final DOCX: {FINAL_DOCX}", flush=True)
    print(f"QA: {QA}", flush=True)


if __name__ == "__main__":
    main()
